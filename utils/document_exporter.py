from docx import Document

def export_to_word(report):

    doc = Document()
    doc.add_heading("Incident Investigation Report", level=1)

    doc.add_paragraph(report)

    file_name = "incident_report.docx"
    doc.save(file_name)

    return file_name
